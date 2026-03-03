from io import BytesIO

from pypdf import PdfReader
from docx import Document as DocxDocument
from odfdo import Document as OdfDocument


extensions = {'.pdf'}
class UnsupportedFileTypeError(Exception):
    pass


def extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    text_parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)
    return ', '.join(text_parts)

'''
TODO: for some reason docx and odfdo dont work. 
'''
def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(BytesIO(file_bytes))
    return ', '.join(para.text for para in doc.paragraphs)


def extract_text_from_odt(file_bytes: bytes) -> str:
    doc = OdfDocument(BytesIO(file_bytes))
    return doc.body.get_formatted_text()


def extract_text(file_bytes: bytes, filename: str = '') -> str:
    fn = filename.lower()
    if fn.endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    elif fn.endswith('.docx'):
        return extract_text_from_docx(file_bytes)
    elif fn.endswith('.odt'):
        return extract_text_from_odt(file_bytes)
    else:
        raise UnsupportedFileTypeError(
            f"Unsupported file type. Supported formats: pdf, docx, odt. Use /text"
        )
